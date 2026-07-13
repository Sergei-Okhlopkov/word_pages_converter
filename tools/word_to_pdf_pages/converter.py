import tempfile
from pathlib import Path

import docx
import fitz  # PyMuPDF
from docx.enum.section import WD_ORIENTATION, WD_SECTION_START
from docx.shared import Inches, Pt
from win32com.client import DispatchEx

_POINTS_PER_INCH = 72.0
_MIN_PAGE_INCHES = 0.5

SUPPORTED_INPUT_EXTENSIONS = {
    ".doc",
    ".docx",
    ".docm",
    ".docb",
    ".dot",
    ".dotx",
    ".dotm",
    ".rtf",
    ".odt",
    ".txt",
    ".htm",
    ".html",
    ".mht",
    ".mhtml",
    ".xml",
}


def ensure_supported_extension(input_path: str) -> None:
    ext = Path(input_path).suffix.lower()
    if ext not in SUPPORTED_INPUT_EXTENSIONS:
        raise ValueError(
            f"Неподдерживаемое расширение: {ext}\n"
            f"Поддерживаются: {', '.join(sorted(SUPPORTED_INPUT_EXTENSIONS))}"
        )


def export_to_pdf_with_word(input_path: str, output_pdf_path: str) -> None:
    """Export the source document to PDF via Microsoft Word."""
    wdFormatPDF = 17

    word = None
    doc = None
    try:
        word = DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(str(Path(input_path).resolve()))
        doc.SaveAs(str(Path(output_pdf_path).resolve()), FileFormat=wdFormatPDF)
    finally:
        if doc is not None:
            try:
                doc.Close(False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass


def _page_size_inches(pdf_page) -> tuple[float, float]:
    page_width = max(_MIN_PAGE_INCHES, pdf_page.rect.width / _POINTS_PER_INCH)
    page_height = max(_MIN_PAGE_INCHES, pdf_page.rect.height / _POINTS_PER_INCH)
    return page_width, page_height


def _configure_section_for_full_page(section, pdf_page) -> tuple[float, float]:
    page_width, page_height = _page_size_inches(pdf_page)
    section.orientation = (
        WD_ORIENTATION.LANDSCAPE if page_width > page_height else WD_ORIENTATION.PORTRAIT
    )
    section.page_width = Inches(page_width)
    section.page_height = Inches(page_height)
    section.left_margin = Inches(0)
    section.right_margin = Inches(0)
    section.top_margin = Inches(0)
    section.bottom_margin = Inches(0)
    section.header_distance = Inches(0)
    section.footer_distance = Inches(0)
    return page_width, page_height


def pdf_to_docx_images(pdf_path: str, output_docx_path: str, dpi: int, on_progress) -> None:
    pdf_doc = fitz.open(pdf_path)
    out_doc = docx.Document()

    try:
        with tempfile.TemporaryDirectory() as temp_dir:
            image_infos = []
            zoom = dpi / 72.0
            matrix = fitz.Matrix(zoom, zoom)
            total_pages = len(pdf_doc)

            for page_index in range(total_pages):
                page = pdf_doc[page_index]
                pix = page.get_pixmap(matrix=matrix, alpha=False)
                img_path = Path(temp_dir) / f"page_{page_index + 1}.png"
                pix.save(str(img_path))
                image_infos.append((img_path, pix.width, pix.height))
                if on_progress:
                    on_progress("render", page_index + 1, total_pages)

            for i, (img_path, img_w_px, img_h_px) in enumerate(image_infos):
                if i > 0:
                    section = out_doc.add_section(WD_SECTION_START.NEW_PAGE)
                else:
                    section = out_doc.sections[0]

                page_width, page_height = _configure_section_for_full_page(section, pdf_doc[i])

                img_aspect = img_w_px / img_h_px
                page_aspect = page_width / page_height

                paragraph = out_doc.add_paragraph()
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                run = paragraph.add_run()

                if img_aspect >= page_aspect:
                    run.add_picture(str(img_path), width=Inches(page_width))
                else:
                    run.add_picture(str(img_path), height=Inches(page_height))

                if on_progress:
                    on_progress("insert", i + 1, len(image_infos))

            out_doc.save(output_docx_path)
    finally:
        pdf_doc.close()


def next_output_path(input_path: str, output_dir: Path | None = None) -> Path:
    input_file = Path(input_path)
    directory = output_dir if output_dir is not None else input_file.parent
    candidate = directory / f"Converted_{input_file.stem}.docx"
    if not candidate.exists():
        return candidate
    return directory / f"Converted_{input_file.stem}1.docx"


def convert_document(input_path: str, on_progress=None, output_dir: Path | None = None) -> Path:
    ensure_supported_extension(input_path)
    output_path = next_output_path(input_path, output_dir=output_dir)
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_pdf = str(Path(temp_dir) / "intermediate.pdf")
        if on_progress:
            on_progress("status", "Экспорт документа в PDF через Microsoft Word...")
        export_to_pdf_with_word(input_path, temp_pdf)
        if on_progress:
            on_progress("status", "Рендер страниц и сборка DOCX...")
        pdf_to_docx_images(temp_pdf, str(output_path), dpi=220, on_progress=on_progress)
    return output_path

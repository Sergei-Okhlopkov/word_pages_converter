import json
import tempfile
import traceback
from pathlib import Path
from threading import Thread
from tkinter import DoubleVar, StringVar, Tk, filedialog
from tkinter import ttk

import docx
import fitz  # PyMuPDF
import pythoncom
from docx.enum.section import WD_SECTION_START
from docx.shared import Inches, Pt
from win32com.client import DispatchEx


SUPPORTED_INPUT_EXTENSIONS = {
    ".doc",
    ".docx",
    ".docm",
    ".docb",
    ".dot",
    ".dotx",
    ".dotm",
    ".rtf",
    ".odt",
    ".txt",
    ".htm",
    ".html",
    ".mht",
    ".mhtml",
    ".xml",
}

THEMES = {
    "light": {
        "bg": "#f5f7fb",
        "panel": "#ffffff",
        "text": "#1f2937",
        "muted": "#6b7280",
        "accent": "#2563eb",
        "entry_bg": "#ffffff",
        "entry_fg": "#111827",
        "log_bg": "#f9fafb",
        "log_fg": "#111827",
    },
    "dark": {
        "bg": "#111827",
        "panel": "#1f2937",
        "text": "#e5e7eb",
        "muted": "#9ca3af",
        "accent": "#3b82f6",
        "entry_bg": "#0f172a",
        "entry_fg": "#f3f4f6",
        "log_bg": "#0b1220",
        "log_fg": "#e5e7eb",
    },
}

STATUS_ICONS = {
    "idle": "◌",
    "working": "⏳",
    "success": "✅",
    "error": "❌",
}


def get_settings_path() -> Path:
    settings_dir = Path.home() / "AppData" / "Local" / "WordPagesToImages"
    settings_dir.mkdir(parents=True, exist_ok=True)
    return settings_dir / "settings.json"


def load_settings() -> dict:
    settings_path = get_settings_path()
    if not settings_path.exists():
        return {"theme": "light"}
    try:
        return json.loads(settings_path.read_text(encoding="utf-8"))
    except Exception:
        return {"theme": "light"}


def save_settings(settings: dict) -> None:
    settings_path = get_settings_path()
    settings_path.write_text(json.dumps(settings, ensure_ascii=False, indent=2), encoding="utf-8")


def ensure_supported_extension(input_path: str) -> None:
    ext = Path(input_path).suffix.lower()
    if ext not in SUPPORTED_INPUT_EXTENSIONS:
        raise ValueError(
            f"Неподдерживаемое расширение: {ext}\n"
            f"Поддерживаются: {', '.join(sorted(SUPPORTED_INPUT_EXTENSIONS))}"
        )


def _default_margins_inches() -> dict:
    """Fallback if PageSetup cannot be read (matches typical Word 'Normal')."""
    return {
        "left": 1.0,
        "right": 1.0,
        "top": 1.0,
        "bottom": 1.0,
        "header": 0.5,
        "footer": 0.5,
    }


def export_to_pdf_with_word(input_path: str, output_pdf_path: str) -> dict:
    """Export via Word; returns page margins from the source document (inches)."""
    wdFormatPDF = 17

    word = None
    doc = None
    margins = _default_margins_inches()
    try:
        word = DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(str(Path(input_path).resolve()))
        ps = doc.PageSetup
        # Word PageSetup distances are in points; 72 pt = 1 inch
        margins = {
            "left": float(ps.LeftMargin) / 72.0,
            "right": float(ps.RightMargin) / 72.0,
            "top": float(ps.TopMargin) / 72.0,
            "bottom": float(ps.BottomMargin) / 72.0,
            "header": float(ps.HeaderDistance) / 72.0,
            "footer": float(ps.FooterDistance) / 72.0,
        }
        doc.SaveAs(str(Path(output_pdf_path).resolve()), FileFormat=wdFormatPDF)
    finally:
        if doc is not None:
            try:
                doc.Close(False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
    return margins


def configure_section(section, margins: dict) -> None:
    section.left_margin = Inches(margins["left"])
    section.right_margin = Inches(margins["right"])
    section.top_margin = Inches(margins["top"])
    section.bottom_margin = Inches(margins["bottom"])
    section.header_distance = Inches(margins["header"])
    section.footer_distance = Inches(margins["footer"])


def pdf_to_docx_images(pdf_path: str, output_docx_path: str, dpi: int, on_progress, margins: dict) -> None:
    pdf_doc = fitz.open(pdf_path)
    out_doc = docx.Document()
    configure_section(out_doc.sections[0], margins)

    try:
        with tempfile.TemporaryDirectory() as temp_dir:
            image_infos = []
            zoom = dpi / 72.0
            matrix = fitz.Matrix(zoom, zoom)
            total_pages = len(pdf_doc)

            for page_index in range(total_pages):
                page = pdf_doc[page_index]
                pix = page.get_pixmap(matrix=matrix, alpha=False)
                img_path = Path(temp_dir) / f"page_{page_index + 1}.png"
                pix.save(str(img_path))
                image_infos.append((img_path, pix.width, pix.height))
                if on_progress:
                    on_progress("render", page_index + 1, total_pages)

            for i, (img_path, img_w_px, img_h_px) in enumerate(image_infos):
                if i > 0:
                    section = out_doc.add_section(WD_SECTION_START.NEW_PAGE)
                    configure_section(section, margins)
                else:
                    section = out_doc.sections[0]

                max_width_inches = section.page_width.inches - section.left_margin.inches - section.right_margin.inches
                max_height_inches = section.page_height.inches - section.top_margin.inches - section.bottom_margin.inches
                # Keep tiny reserve to prevent Word from moving image to a new page.
                max_height_inches = max(0.5, max_height_inches - 0.02)

                img_aspect = img_w_px / img_h_px
                box_aspect = max_width_inches / max_height_inches

                paragraph = out_doc.add_paragraph()
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                run = paragraph.add_run()

                if img_aspect >= box_aspect:
                    run.add_picture(str(img_path), width=Inches(max_width_inches))
                else:
                    run.add_picture(str(img_path), height=Inches(max_height_inches))

                if on_progress:
                    on_progress("insert", i + 1, len(image_infos))

            out_doc.save(output_docx_path)
    finally:
        pdf_doc.close()


def next_output_path(input_path: str) -> Path:
    input_file = Path(input_path)
    base = input_file.with_name(f"Converted_{input_file.stem}")
    candidate = base.with_suffix(".docx")
    if not candidate.exists():
        return candidate
    # As requested: when name exists, append digit 1 at end.
    return input_file.with_name(f"{base.name}1.docx")


def convert_document(input_path: str, on_progress=None) -> Path:
    ensure_supported_extension(input_path)
    output_path = next_output_path(input_path)
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_pdf = str(Path(temp_dir) / "intermediate.pdf")
        if on_progress:
            on_progress("status", "Экспорт документа в PDF через Microsoft Word...")
        margins = export_to_pdf_with_word(input_path, temp_pdf)
        if on_progress:
            on_progress("status", "Рендер страниц и сборка DOCX...")
        pdf_to_docx_images(temp_pdf, str(output_path), dpi=220, on_progress=on_progress, margins=margins)
    return output_path


class App:
    def __init__(self, root: Tk):
        self.root = root
        self.root.title("Word Pages Converter")
        self.root.geometry("900x620")
        self.root.minsize(900, 620)

        self.settings = load_settings()
        self.theme_name = self.settings.get("theme", "light")
        if self.theme_name not in THEMES:
            self.theme_name = "light"

        self.input_var = StringVar()
        self.status_var = StringVar(value="Ожидание запуска")
        self.status_icon_var = StringVar(value=STATUS_ICONS["idle"])
        self.progress_var = DoubleVar(value=0.0)
        self.is_busy = False

        self.style = ttk.Style()
        self._build_ui()
        self.apply_theme()

    def _build_ui(self):
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)

        self.main = ttk.Frame(self.root, padding=16)
        self.main.grid(row=0, column=0, sticky="nsew")
        self.main.columnconfigure(0, weight=1)
        self.main.rowconfigure(3, weight=1)

        top = ttk.Frame(self.main)
        top.grid(row=0, column=0, sticky="ew", pady=(0, 12))
        top.columnconfigure(3, weight=1)

        self.theme_btn = ttk.Button(top, text="Тема: ...", command=self.toggle_theme, width=14)
        self.theme_btn.grid(row=0, column=0, sticky="w", padx=(0, 10))

        ttk.Label(top, text="Статус").grid(row=0, column=1, sticky="w", padx=(0, 6))
        self.status_icon_lbl = ttk.Label(top, textvariable=self.status_icon_var, font=("Segoe UI Emoji", 12))
        self.status_icon_lbl.grid(row=0, column=2, sticky="w", padx=(0, 8))
        self.status_lbl = ttk.Label(top, textvariable=self.status_var)
        self.status_lbl.grid(row=0, column=3, sticky="ew")

        card = ttk.Frame(self.main, padding=16)
        card.grid(row=1, column=0, sticky="ew", pady=(0, 12))
        card.columnconfigure(1, weight=1)

        ttk.Label(card, text="Входной Word-файл:").grid(row=0, column=0, sticky="w")
        self.input_entry = ttk.Entry(card, textvariable=self.input_var)
        self.input_entry.grid(row=1, column=0, columnspan=2, sticky="ew", pady=(8, 0))
        self.pick_btn = ttk.Button(card, text="Выбрать файл", command=self.pick_input)
        self.pick_btn.grid(row=1, column=2, sticky="e", padx=(8, 0))

        self.run_btn = ttk.Button(card, text="Конвертировать", command=self.on_convert)
        self.run_btn.grid(row=2, column=0, sticky="w", pady=(12, 0))

        self.progress = ttk.Progressbar(self.main, variable=self.progress_var, maximum=100, mode="determinate")
        self.progress.grid(row=2, column=0, sticky="ew", pady=(0, 12))

        legend_card = ttk.Frame(self.main, padding=12)
        legend_card.grid(row=3, column=0, sticky="nsew")
        legend_card.columnconfigure(0, weight=1)
        ttk.Label(legend_card, text="Легенда иконок").grid(row=0, column=0, sticky="w", pady=(0, 8))
        ttk.Label(
            legend_card,
            text="◌ ожидание    ⏳ идёт работа    ✅ успех    ❌ ошибка",
            font=("Segoe UI", 10),
        ).grid(row=1, column=0, sticky="w")

    def apply_theme(self):
        colors = THEMES[self.theme_name]
        self.root.configure(bg=colors["bg"])
        self.style.theme_use("clam")

        self.style.configure(".", background=colors["bg"], foreground=colors["text"])
        self.style.configure("TFrame", background=colors["bg"])
        self.style.configure("TLabel", background=colors["bg"], foreground=colors["text"], font=("Segoe UI", 10))
        self.style.configure(
            "TButton",
            background=colors["accent"],
            foreground="#ffffff",
            padding=(10, 8),
            borderwidth=1,
            relief="flat",
            font=("Segoe UI", 10, "bold"),
        )
        self.style.map(
            "TButton",
            background=[("active", colors["accent"])],
            foreground=[("disabled", "#9ca3af"), ("active", "#ffffff")],
        )
        self.style.configure(
            "TEntry",
            fieldbackground=colors["entry_bg"],
            foreground=colors["entry_fg"],
            bordercolor=colors["muted"],
            lightcolor=colors["muted"],
            darkcolor=colors["muted"],
            padding=6,
        )
        self.style.configure(
            "Horizontal.TProgressbar",
            background=colors["accent"],
            troughcolor=colors["panel"],
            bordercolor=colors["panel"],
            lightcolor=colors["accent"],
            darkcolor=colors["accent"],
        )

        self.main.configure(style="TFrame")
        self.theme_btn.configure(text=f"Тема: {'Тёмная' if self.theme_name == 'dark' else 'Светлая'}")
    def set_status(self, kind: str, text: str):
        self.status_icon_var.set(STATUS_ICONS.get(kind, STATUS_ICONS["idle"]))
        self.status_var.set(text)

    def toggle_theme(self):
        self.theme_name = "dark" if self.theme_name == "light" else "light"
        self.settings["theme"] = self.theme_name
        save_settings(self.settings)
        self.apply_theme()
        self.set_status("idle", f"Тема: {'тёмная' if self.theme_name == 'dark' else 'светлая'}")

    def set_busy(self, busy: bool):
        self.is_busy = busy
        state = "disabled" if busy else "normal"
        self.run_btn.configure(state=state)
        self.pick_btn.configure(state=state)
        self.input_entry.configure(state=state)

    def pick_input(self):
        filetypes = [
            (
                "Word and text documents",
                "*.doc *.docx *.docm *.docb *.dot *.dotx *.dotm *.rtf *.odt *.txt *.htm *.html *.mht *.mhtml *.xml",
            )
        ]
        selected = filedialog.askopenfilename(title="Выберите входной файл", filetypes=filetypes)
        if selected:
            self.input_var.set(selected)
            self.set_status("idle", "Файл выбран. Нажмите 'Конвертировать'.")

    def on_convert(self):
        if self.is_busy:
            return
        input_path = self.input_var.get().strip()
        if not input_path:
            self.set_status("error", "Не выбран входной файл.")
            return

        self.progress_var.set(0)
        self.set_busy(True)
        self.set_status("working", "Подготовка конвертации...")

        worker = Thread(target=self._convert_worker, args=(input_path,), daemon=True)
        worker.start()

    def _ui_progress(self, stage: str, current=None, total=None):
        if stage == "status":
            self.set_status("working", str(current))
            return

        if not total:
            return
        total = max(1, int(total))
        current = int(current)

        if stage == "render":
            percent = 10 + (current / total) * 45
            self.set_status("working", f"Рендер страниц: {current}/{total}")
        elif stage == "insert":
            percent = 55 + (current / total) * 44
            self.set_status("working", f"Сборка DOCX: {current}/{total}")
        else:
            return

        self.progress_var.set(min(99, percent))

    def _convert_worker(self, input_path: str):
        def callback(*args):
            self.root.after(0, self._ui_progress, *args)

        try:
            pythoncom.CoInitialize()
            output = convert_document(input_path, on_progress=callback)
            self.root.after(0, self._on_success, str(output))
        except Exception as exc:
            details = "".join(traceback.format_exception_only(type(exc), exc)).strip()
            self.root.after(0, self._on_error, details)
        finally:
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass

    def _on_success(self, output_path: str):
        self.progress_var.set(100)
        self.set_status("success", f"Готово: {output_path}")
        self.set_busy(False)

    def _on_error(self, details: str):
        self.progress_var.set(0)
        self.set_status("error", f"Ошибка: {details}")
        self.set_busy(False)


def main():
    root = Tk()
    App(root)
    root.mainloop()


if __name__ == "__main__":
    main()
